from __future__ import annotations
import io
import os
import platform
import time
from pathlib import Path
from typing import Any

                                                              


def _safe_resolve_under(base: Path, child: Path) -> Path | None:
    try:
        rbase = base.resolve()
        r = child.resolve()
        if r == rbase or rbase in r.parents:
            return r if r.is_file() else None
    except OSError:
        return None
    return None


def url_to_local_file(base_dir: Path, url: str | None) -> Path | None:
    if not url or not isinstance(url, str):
        return None
    u = url.strip().split("?")[0]
    if not u.startswith("/"):
        return None
    name = Path(u).name
    if not name or name != Path(u.replace("\\", "/")).name:
        return None
    if u.startswith("/uploads/"):
        return _safe_resolve_under(base_dir, base_dir / "uploads" / name)
    if u.startswith("/results/"):
        return _safe_resolve_under(base_dir, base_dir / "results" / name)
    if u.startswith("/static/"):
        rel = u[len("/static/") :].lstrip("/").replace("\\", "/")
        if not rel or ".." in rel or rel.startswith("/"):
            return None
        parts = Path(rel).parts
        if ".." in parts:
            return None
        return _safe_resolve_under(base_dir, base_dir / "static" / rel)
    return None


def _img_ext_ok(p: Path) -> bool:
    return p.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp", ".bmp"}


def _xml(s: Any) -> str:
    return str(s if s is not None else "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


                                                                       


def build_markdown(payload: dict[str, Any]) -> str:
    title = str(payload.get("title") or "桥梁表观病害检测报告")
    lines: list[str] = [f"# {title}", ""]
    meta = payload.get("meta") or {}
    lines += ["## 元信息", ""]
    for k, v in [
        ("检测时间", meta.get("time")),
        ("使用模型", meta.get("model")),
        ("病害总数", meta.get("total")),
        ("平均置信度", meta.get("conf")),
        ("推理耗时", meta.get("ms")),
    ]:
        if v:
            lines.append(f"- **{k}**：{v}")
    lines.append("")

    lines += ["## 一、检测概要", "", str(payload.get("summary_lead_plain") or "").strip(), ""]
    for card in payload.get("summary_cards") or []:
        idx = card.get("index", 0)
        lines.append(f"### 记录 {idx}")
        lines.append(f"- 时间：{card.get('time', '')}")
        lines.append(f"- 模型：{card.get('model', '')}")
        lines.append(f"- 病害数：{card.get('total', '')}（JSON 明细 {card.get('n_dets', 0)} 条）")
        lines.append(f"- 平均置信度：{card.get('avg_conf', '')}")
        lines.append(f"- 推理耗时：{card.get('infer_ms', '')}")
        cats = card.get("top_cats") or []
        if cats:
            lines.append("- 主要类别：" + "；".join(f"{c.get('name')} {c.get('count')} 处" for c in cats))
        lines.append("")

    lines += ["## 二、病害分类统计（合并）", ""]
    ms = payload.get("merged_stats") or {}
    if isinstance(ms, dict) and ms:
        for k, v in sorted(ms.items(), key=lambda kv: (-int(kv[1] or 0), str(kv[0]))):
            lines.append(f"- **{k}**：{v} 处")
    else:
        lines.append("- （无）")
    lines.append("")

    lines += ["## 三、现场图像", ""]
    for fig in payload.get("figures") or []:
        lines.append(f"### {fig.get('title', '配图')}")
        ou = fig.get("original_url") or ""
        ru = fig.get("result_url") or ""
        lines.append(f"- 原图：<{ou}>" if ou else "- 原图：（无）")
        lines.append(f"- 结果：<{ru}>" if ru else "- 结果：（无）")
        lines.append("")

    lines += ["## 四、病害明细", ""]
    rows = payload.get("detail_rows") or []
    if rows:
        lines.append("| 序号 | 来源 | 类型 | 置信度 | 备注 |")
        lines.append("| --- | --- | --- | --- | --- |")
        for r in rows:
            def esc(s: Any) -> str:
                t = str(s if s is not None else "").replace("|", "\\|").replace("\n", " ")
                return t

            lines.append(
                f"| {esc(r.get('seq'))} | {esc(r.get('source'))} | {esc(r.get('class'))} | {esc(r.get('conf'))} | {esc(r.get('remark'))} |"
            )
    else:
        lines.append("（无）")
    lines.append("")

    fl = payload.get("footer_line") or ""
    fg = payload.get("footer_generated") or ""
    if fl or fg:
        lines += ["---", "", fl, fg, ""]
    return "\n".join(lines)


                                                                       


def build_docx(payload: dict[str, Any], base_dir: Path) -> bytes:
    from docx import Document                
    from docx.enum.text import WD_ALIGN_PARAGRAPH                
    from docx.shared import Cm, Pt                

    doc = Document()
    title = str(payload.get("title") or "桥梁表观病害检测报告")
    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("元信息", level=1)
    meta = payload.get("meta") or {}
    for label, key in [
        ("检测时间", "time"),
        ("使用模型", "model"),
        ("病害总数", "total"),
        ("平均置信度", "conf"),
        ("推理耗时", "ms"),
    ]:
        v = meta.get(key)
        if v:
            doc.add_paragraph(f"{label}：{v}", style=None)

    doc.add_heading("一、检测概要", level=1)
    doc.add_paragraph(str(payload.get("summary_lead_plain") or "").strip())

    for card in payload.get("summary_cards") or []:
        doc.add_heading(f"记录 {card.get('index', '')}", level=2)
        doc.add_paragraph(f"时间：{card.get('time', '')}")
        doc.add_paragraph(f"模型：{card.get('model', '')}")
        doc.add_paragraph(f"病害数：{card.get('total', '')}（JSON 明细 {card.get('n_dets', 0)} 条）")
        doc.add_paragraph(f"平均置信度：{card.get('avg_conf', '')}")
        doc.add_paragraph(f"推理耗时：{card.get('infer_ms', '')}")
        cats = card.get("top_cats") or []
        if cats:
            doc.add_paragraph("主要类别：" + "；".join(f"{c.get('name')} {c.get('count')} 处" for c in cats))

    doc.add_heading("二、病害分类统计（合并）", level=1)
    ms = payload.get("merged_stats") or {}
    if isinstance(ms, dict) and ms:
        for k, v in sorted(ms.items(), key=lambda kv: (-int(kv[1] or 0), str(kv[0]))):
            doc.add_paragraph(f"{k}：{v} 处")
    else:
        doc.add_paragraph("（无）")

    doc.add_heading("三、现场图像", level=1)
    for fig in payload.get("figures") or []:
        doc.add_heading(str(fig.get("title") or "配图"), level=2)
        ou = url_to_local_file(base_dir, fig.get("original_url"))
        ru = url_to_local_file(base_dir, fig.get("result_url"))
        if ou and _img_ext_ok(ou):
            doc.add_paragraph("原始图像")
            doc.add_picture(str(ou), width=Cm(7.5))
        elif fig.get("original_url"):
            doc.add_paragraph(f"原始图像（未嵌入，路径或格式不支持）：{fig.get('original_url')}")
        if ru and _img_ext_ok(ru):
            doc.add_paragraph("检测结果")
            doc.add_picture(str(ru), width=Cm(7.5))
        elif fig.get("result_url"):
            doc.add_paragraph(f"检测结果（未嵌入，路径或格式不支持）：{fig.get('result_url')}")

    doc.add_heading("四、病害明细", level=1)
    rows = payload.get("detail_rows") or []
    if rows:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "序号"
        hdr[1].text = "来源"
        hdr[2].text = "病害类型"
        hdr[3].text = "置信度"
        hdr[4].text = "备注"
        for r in rows:
            row = table.add_row().cells
            row[0].text = str(r.get("seq", ""))
            row[1].text = str(r.get("source", ""))
            row[2].text = str(r.get("class", ""))
            row[3].text = str(r.get("conf", ""))
            row[4].text = str(r.get("remark", ""))
        for row in table.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
    else:
        doc.add_paragraph("（无）")

    fl = payload.get("footer_line") or ""
    fg = payload.get("footer_generated") or ""
    if fl or fg:
        doc.add_paragraph("")
        doc.add_paragraph(fl)
        doc.add_paragraph(fg)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


                                                         

_cn_font_registered: str | None = None

                                                        
_LINUX_MACOS_CN_FONT_PATHS: tuple[str, ...] = (
                                                                             
    "/usr/share/fonts/google-noto-cjk/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/google-noto-cjk/NotoSansCJKsc-Regular.otf",
                                                        
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc",
         
    "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
    "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
                
    "/usr/share/fonts/truetype/arphic/uming.ttc",
    "/usr/share/fonts/truetype/arphic/ukai.ttc",
              
    "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
                                    
    "/Library/Fonts/Arial Unicode.ttf",
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
)


def _pdf_cn_font_candidate_paths() -> list[Path]:
                                                          
    out: list[Path] = []
    seen: set[str] = set()

    def add(p: Path) -> None:
        try:
            k = str(p.resolve())
        except OSError:
            k = str(p)
        if k not in seen:
            seen.add(k)
            out.append(p)

    env = os.environ.get("REPORT_EXPORT_FONT", "").strip()
    if env:
        add(Path(env))

    pkg_fonts = Path(__file__).resolve().parent / "fonts"
    for n in (
        "NotoSansSC-Regular.otf",
        "NotoSansSC-VF.ttf",
        "SourceHanSansSC-Regular.otf",
        "SourceHanSansSC-VF.ttf",
        "msyh.ttf",
        "msyh.ttc",
        "simhei.ttf",
        "simsun.ttc",
    ):
        add(pkg_fonts / n)

    if platform.system() == "Windows":
        wind = Path(os.environ.get("WINDIR", r"C:\Windows"))
        fonts = wind / "Fonts"
        for n in ("msyh.ttf", "msyh.ttc", "simhei.ttf", "simsun.ttc", "simsunb.ttf"):
            add(fonts / n)
    else:
        for s in _LINUX_MACOS_CN_FONT_PATHS:
            add(Path(s))

    return out


def _ensure_pdf_cn_font() -> str:
    global _cn_font_registered
    if _cn_font_registered:
        return _cn_font_registered
    from reportlab.pdfbase import pdfmetrics                
    from reportlab.pdfbase.ttfonts import TTFont                

    for p in _pdf_cn_font_candidate_paths():
        try:
            if p.is_file():
                pdfmetrics.registerFont(TTFont("ReportCN", str(p)))
                _cn_font_registered = "ReportCN"
                return "ReportCN"
        except Exception:
            continue
    raise ValueError(
        "无法为 PDF 注册中文字体。部署服务器请任选其一："
        "① 安装系统字体——Debian/Ubuntu：`sudo apt-get update && sudo apt-get install -y fonts-noto-cjk`（或 `fonts-wqy-zenhei`）；"
        "OpenCloudOS / CentOS / RHEL 等无 apt，请用：`sudo dnf install -y google-noto-sans-cjk-ttc-fonts`（若无 dnf 则用 `yum` 代替 `dnf`），或 `sudo yum install -y wqy-microhei-fonts`；"
        "② 设置环境变量 `REPORT_EXPORT_FONT` 为服务器上某个 .ttf/.ttc/.otf 的绝对路径；"
        "③ 将字体文件放到本仓库目录 `fonts/` 下（如从 https://fonts.google.com/noto/specimen/Noto+Sans+SC 下载 Regular 放入 `fonts/NotoSansSC-Regular.otf`）。"
    )


def build_pdf(payload: dict[str, Any], base_dir: Path) -> bytes:
    from reportlab.lib import colors                
    from reportlab.lib.pagesizes import A4                
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet                
    from reportlab.lib.units import cm                
    from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle                

    font = _ensure_pdf_cn_font()
    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "CN",
        parent=styles["Normal"],
        fontName=font,
        fontSize=10,
        leading=14,
    )
    h0 = ParagraphStyle("H0", parent=styles["Heading1"], fontName=font, fontSize=18, leading=22, alignment=1)
    h1 = ParagraphStyle("H1", parent=styles["Heading2"], fontName=font, fontSize=14, leading=18)
    h2 = ParagraphStyle("H2", parent=styles["Heading3"], fontName=font, fontSize=12, leading=16)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=2 * cm, leftMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm)

    story: list[Any] = []
    title = str(payload.get("title") or "桥梁表观病害检测报告")
    story.append(Paragraph(_xml(title), h0))
    story.append(Spacer(1, 0.4 * cm))

    meta = payload.get("meta") or {}
    for label, key in [
        ("检测时间", "time"),
        ("使用模型", "model"),
        ("病害总数", "total"),
        ("平均置信度", "conf"),
        ("推理耗时", "ms"),
    ]:
        v = meta.get(key)
        if v:
            story.append(Paragraph(_xml(f"{label}：{v}"), normal))
    story.append(Spacer(1, 0.3 * cm))

    story.append(Paragraph("一、检测概要", h1))
    lead = str(payload.get("summary_lead_plain") or "")
    story.append(Paragraph(_xml(lead).replace("\n", "<br/>"), normal))
    story.append(Spacer(1, 0.2 * cm))

    for card in payload.get("summary_cards") or []:
        story.append(Paragraph(_xml(f"记录 {card.get('index', '')}"), h2))
        for line in [
            f"时间：{card.get('time', '')}",
            f"模型：{card.get('model', '')}",
            f"病害数：{card.get('total', '')}（JSON {card.get('n_dets', 0)} 条）",
            f"平均置信度：{card.get('avg_conf', '')}",
            f"推理耗时：{card.get('infer_ms', '')}",
        ]:
            story.append(Paragraph(_xml(line), normal))
        cats = card.get("top_cats") or []
        if cats:
            t = "主要类别：" + "；".join(f"{c.get('name')} {c.get('count')} 处" for c in cats)
            story.append(Paragraph(_xml(t), normal))
        story.append(Spacer(1, 0.15 * cm))

    story.append(Paragraph("二、病害分类统计（合并）", h1))
    ms = payload.get("merged_stats") or {}
    if isinstance(ms, dict) and ms:
        for k, v in sorted(ms.items(), key=lambda kv: (-int(kv[1] or 0), str(kv[0]))):
            story.append(Paragraph(_xml(f"{k}：{v} 处"), normal))
    else:
        story.append(Paragraph("（无）", normal))
    story.append(Spacer(1, 0.3 * cm))

    story.append(Paragraph("三、现场图像", h1))
    for fig in payload.get("figures") or []:
        story.append(Paragraph(_xml(str(fig.get("title") or "配图")), h2))
        ou = url_to_local_file(base_dir, fig.get("original_url"))
        ru = url_to_local_file(base_dir, fig.get("result_url"))
        for label, p in [("原始", ou), ("结果", ru)]:
            if p and _img_ext_ok(p):
                try:
                    story.append(Paragraph(_xml(label), normal))
                    w = 8 * cm
                    story.append(Image(str(p), width=w, height=None))
                except Exception:
                    story.append(Paragraph(_xml(f"{label}图加载失败：{p.name}"), normal))
            elif fig.get("original_url") or fig.get("result_url"):
                u = fig.get("original_url") if label == "原始" else fig.get("result_url")
                story.append(Paragraph(_xml(f"{label}：（未嵌入）{u}"), normal))
        story.append(Spacer(1, 0.2 * cm))

    story.append(Paragraph("四、病害明细", h1))
    rows = payload.get("detail_rows") or []
    if rows:
        data = [["序号", "来源", "类型", "置信度", "备注"]]
        for r in rows:
            data.append(
                [
                    _xml(str(r.get("seq", ""))),
                    _xml(str(r.get("source", ""))[:24]),
                    _xml(str(r.get("class", ""))[:20]),
                    _xml(str(r.get("conf", ""))),
                    _xml(str(r.get("remark", ""))[:40]),
                ]
            )
        t = Table(data, colWidths=[1.2 * cm, 2.8 * cm, 2.5 * cm, 2 * cm, 5.5 * cm])
        t.setStyle(
            TableStyle(
                [
                    ("FONT", (0, 0), (-1, -1), font, 8),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(t)
    else:
        story.append(Paragraph("（无）", normal))

    fl = payload.get("footer_line") or ""
    fg = payload.get("footer_generated") or ""
    if fl or fg:
        story.append(Spacer(1, 0.4 * cm))
        story.append(Paragraph(_xml(fl), normal))
        story.append(Paragraph(_xml(fg), normal))

    doc.build(story)
    return buf.getvalue()


def build_report_file(fmt: str, payload: dict[str, Any], base_dir: Path) -> tuple[bytes, str, str]:
    fmt = fmt.strip().lower()
    stamp = time.strftime("%Y%m%d_%H%M%S")
    if fmt == "md":
        body = build_markdown(payload).encode("utf-8")
        return body, f"桥梁检测报告_{stamp}.md", "text/markdown; charset=utf-8"
    if fmt == "docx":
        body = build_docx(payload, base_dir)
        return body, f"桥梁检测报告_{stamp}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if fmt == "pdf":
        body = build_pdf(payload, base_dir)
        return body, f"桥梁检测报告_{stamp}.pdf", "application/pdf"
    raise ValueError("不支持的格式")
